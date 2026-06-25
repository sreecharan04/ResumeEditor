import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_section_header(doc: Document, title: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "404040")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_right_tab(para, pos_twips: int = 8640):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def generate_docx(resume_data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    # ── Name ──────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(resume_data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Calibri"

    # ── Contact ───────────────────────────────────────────────────────────
    contact = resume_data.get("contact", {})
    parts = [
        v
        for k in ("email", "phone", "location", "linkedin", "github", "website")
        if (v := contact.get(k))
    ]
    if parts:
        cp = doc.add_paragraph(" | ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(6)
        for run in cp.runs:
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    # ── Summary ───────────────────────────────────────────────────────────
    if resume_data.get("summary"):
        _add_section_header(doc, "Summary")
        p = doc.add_paragraph(resume_data["summary"])
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # ── Experience ────────────────────────────────────────────────────────
    if resume_data.get("experience"):
        _add_section_header(doc, "Experience")
        for exp in resume_data["experience"]:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(1)
            _add_right_tab(para)

            title_text = exp.get("company", "")
            if exp.get("title"):
                title_text += f"  ·  {exp['title']}"
            if exp.get("location"):
                title_text += f",  {exp['location']}"
            r1 = para.add_run(title_text)
            r1.bold = True
            r1.font.size = Pt(10.5)
            r1.font.name = "Calibri"

            if exp.get("dates"):
                para.add_run("\t")
                r2 = para.add_run(exp["dates"])
                r2.italic = True
                r2.font.size = Pt(10)
                r2.font.name = "Calibri"

            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.15)
                br = bp.add_run(bullet)
                br.font.size = Pt(10)
                br.font.name = "Calibri"

    # ── Education ─────────────────────────────────────────────────────────
    if resume_data.get("education"):
        _add_section_header(doc, "Education")
        for edu in resume_data["education"]:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(1)
            _add_right_tab(para)

            degree_text = edu.get("institution", "")
            if edu.get("degree"):
                degree_text += f"  ·  {edu['degree']}"
                if edu.get("field"):
                    degree_text += f" in {edu['field']}"
            if edu.get("gpa"):
                degree_text += f"  ·  GPA: {edu['gpa']}"
            r1 = para.add_run(degree_text)
            r1.bold = True
            r1.font.size = Pt(10.5)
            r1.font.name = "Calibri"

            if edu.get("dates"):
                para.add_run("\t")
                r2 = para.add_run(edu["dates"])
                r2.italic = True
                r2.font.size = Pt(10)
                r2.font.name = "Calibri"

            for detail in edu.get("details", []):
                dp = doc.add_paragraph(detail)
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after = Pt(1)
                for run in dp.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    # ── Skills ────────────────────────────────────────────────────────────
    if resume_data.get("skills"):
        _add_section_header(doc, "Skills")
        for skill_section in resume_data["skills"]:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(1)
            if skill_section.get("category"):
                cat_run = para.add_run(f"{skill_section['category']}: ")
                cat_run.bold = True
                cat_run.font.size = Pt(10)
                cat_run.font.name = "Calibri"
            items_run = para.add_run(", ".join(skill_section.get("items", [])))
            items_run.font.size = Pt(10)
            items_run.font.name = "Calibri"

    # ── Projects ──────────────────────────────────────────────────────────
    if resume_data.get("projects"):
        _add_section_header(doc, "Projects")
        for proj in resume_data["projects"]:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(1)
            nr = para.add_run(proj.get("name", ""))
            nr.bold = True
            nr.font.size = Pt(10.5)
            nr.font.name = "Calibri"
            if proj.get("technologies"):
                tr = para.add_run(f"  ·  {proj['technologies']}")
                tr.italic = True
                tr.font.size = Pt(10)
                tr.font.name = "Calibri"

            if proj.get("description"):
                dp = doc.add_paragraph(proj["description"])
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after = Pt(1)
                for run in dp.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

            for bullet in proj.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.15)
                br = bp.add_run(bullet)
                br.font.size = Pt(10)
                br.font.name = "Calibri"

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
